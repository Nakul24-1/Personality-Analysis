
from docx import *
import re
import json

#----------01_Import File Name----------
document = Document('projectlink.docx')  #Change filename here
#02_-----------Declare Variables-----------
bolds=[]
emails=[]
phones=[]
#-----------03_Extract Elements From the Word File-----------
for para in document.paragraphs:
 
    #03.1 Find email and phone numbers within the paragraph text
    text = para.text
    print(text)
    """table = document.tables[0]
    for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print (paragraph.text)"""
    
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
    email_list = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+",text)
    phone_list=re.findall(r'[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]',text)
 
    for email in email_list:
        emails.append(email)
 
    for phone in phone_list:
        phones.append(phone)
    for i in para.text:
        print(i,end="")
 
    #03.2 Find the bold style within the word document
    for run in para.runs:
        if run.bold :
            bolds.append(run.text)
11
#-----------04_Create Output-----------
style_Dict={'emails':emails,
              'phone_numbers':phones,
              'bold_phrases':bolds
              }
 
print("\nWord File Output:\n")

with open ('cv.txt','w',encoding='utf-8') as f:
    for para in document.paragraphs:
        text = para.text
        f.write(text)
    for i in para.text:
        f.write(i,end="")
    
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    f.write(paragraph.text)
                    f.write(" ")
    
    
r = json.dumps(style_Dict)
loaded_r = json.loads(r)

print("\n",json.dumps(loaded_r,indent=4, sort_keys=False))  #Pretty print the JSON output
